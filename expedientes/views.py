import os
import re
import json
import zipfile
import io
from datetime import datetime
from pathlib import Path

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, logout
from django.http import JsonResponse, HttpResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_http_methods
from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage

from .models import Estudiante, Expediente1, Expediente2, Resolucion, ESCUELAS


# ──────────────────────────────────────────────────────────────
# HELPERS DE RENOMBRADO
# ──────────────────────────────────────────────────────────────

def renombrar_exp1(dni, apellidos_nombres, codigo_matricula=''):
    """Retorna dict con los nombres correctos para cada archivo del expediente 1"""
    dni = (dni or '').strip()
    nombre = (apellidos_nombres or '').strip()
    codigo = (codigo_matricula or '').strip()
    foto_nombre = f'{codigo}_{nombre}.jpg' if codigo else f'{nombre}.jpg'
    return {
        'dni_pdf':               f'R05_{dni}_DNI.pdf',
        'solicitud_pdf':         f'R04_{dni}_SOL.pdf',
        'recibo_pdf':            f'R03_{dni}_PAG.pdf',
        'registro_graduados_pdf':f'R01_{dni}_RSC.pdf',
        'foto':                  foto_nombre,
    }


def renombrar_exp2(apellidos_nombres):
    """Retorna dict con los nombres correctos para cada archivo del expediente 2"""
    nombre = (apellidos_nombres or '').strip()
    return {
        'certificado_notas':    f'1 Certificado de Notas {nombre}.pdf',
        'constancia_proyeccion':f'2° Constancia de Proyección Social {nombre}.pdf',
        'constancia_no_adeudo': f'3° Constancia NO Adeudo {nombre}.pdf',
    }


# ──────────────────────────────────────────────────────────────
# AUTENTICACIÓN
# ──────────────────────────────────────────────────────────────

def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    return render(request, 'expedientes/login.html')


# ──────────────────────────────────────────────────────────────
# DASHBOARD - CRUD ESTUDIANTES
# ──────────────────────────────────────────────────────────────

@login_required
def dashboard(request):
    estudiantes = Estudiante.objects.all()
    data = []
    for e in estudiantes:
        archivos = e.archivos_estado()
        total_archivos = sum(archivos['exp1'].values()) + sum(archivos['exp2'].values())
        data.append({
            'obj': e,
            'semaforo': e.estado_semaforo(),
            'archivos': archivos,
            'total_archivos': total_archivos,
        })
    # Ordenar de mayor a menor archivos entregados
    data.sort(key=lambda x: x['total_archivos'], reverse=True)
    return render(request, 'expedientes/dashboard.html', {'estudiantes': data})


@login_required
def crear_estudiante(request):
    """Crea estudiante vacío y redirige al expediente"""
    estudiante = Estudiante.objects.create(creado_por=request.user)
    Expediente1.objects.create(estudiante=estudiante)
    Expediente2.objects.create(estudiante=estudiante)
    return redirect('expediente', pk=estudiante.pk)


@login_required
def eliminar_estudiante(request, pk):
    if request.method == 'POST':
        estudiante = get_object_or_404(Estudiante, pk=pk)
        estudiante.delete()
        return JsonResponse({'ok': True})
    return JsonResponse({'ok': False}, status=405)


# ──────────────────────────────────────────────────────────────
# VISTA PRINCIPAL DE EXPEDIENTE
# ──────────────────────────────────────────────────────────────

@login_required
def expediente(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    exp1, _ = Expediente1.objects.get_or_create(estudiante=estudiante)
    exp2, _ = Expediente2.objects.get_or_create(estudiante=estudiante)
    resoluciones = estudiante.resoluciones.all()

    nombres_exp1 = renombrar_exp1(estudiante.dni, estudiante.apellidos_nombres, estudiante.codigo_matricula)
    nombres_exp2 = renombrar_exp2(estudiante.apellidos_nombres)

    exp1_count = sum([bool(exp1.dni_pdf), bool(exp1.solicitud_pdf), bool(exp1.foto),
                       bool(exp1.recibo_pdf), bool(exp1.registro_graduados_pdf)])
    exp2_count = sum([bool(exp2.certificado_notas), bool(exp2.constancia_proyeccion),
                      bool(exp2.constancia_no_adeudo)])

    return render(request, 'expedientes/expediente.html', {
        'estudiante': estudiante,
        'exp1': exp1,
        'exp2': exp2,
        'escuelas': ESCUELAS,
        'resoluciones': resoluciones,
        'nombres_exp1': nombres_exp1,
        'nombres_exp2': nombres_exp2,
        'exp1_count': exp1_count,
        'exp2_count': exp2_count,
    })


# ──────────────────────────────────────────────────────────────
# GUARDAR DATOS DEL ESTUDIANTE (AJAX)
# ──────────────────────────────────────────────────────────────

@login_required
@require_POST
def guardar_estudiante(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    data = json.loads(request.body)

    campos = ['apellidos_nombres', 'dni', 'n_expediente', 'proveido',
              'escuela', 'celular', 'correo', 'codigo_matricula', 'n_informe']
    for campo in campos:
        if campo in data:
            setattr(estudiante, campo, data[campo])
    estudiante.save()

    # Recalcular nombres
    nombres_exp1 = renombrar_exp1(estudiante.dni, estudiante.apellidos_nombres, estudiante.codigo_matricula)
    nombres_exp2 = renombrar_exp2(estudiante.apellidos_nombres)

    return JsonResponse({
        'ok': True,
        'nombre_display': estudiante.get_nombre_display(),
        'semaforo': estudiante.estado_semaforo(),
        'nombres_exp1': nombres_exp1,
        'nombres_exp2': nombres_exp2,
    })


# ──────────────────────────────────────────────────────────────
# SUBIR ARCHIVO (AJAX)
# ──────────────────────────────────────────────────────────────

@login_required
@require_POST
def subir_archivo(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    expediente_tipo = request.POST.get('expediente')  # 'exp1' o 'exp2'
    campo = request.POST.get('campo')
    archivo = request.FILES.get('archivo')

    if not archivo or not campo or not expediente_tipo:
        return JsonResponse({'ok': False, 'error': 'Datos incompletos'}, status=400)

    if expediente_tipo == 'exp1':
        exp = Expediente1.objects.get_or_create(estudiante=estudiante)[0]
        campos_validos = ['dni_pdf', 'solicitud_pdf', 'foto', 'recibo_pdf', 'registro_graduados_pdf']
        nombres = renombrar_exp1(estudiante.dni, estudiante.apellidos_nombres, estudiante.codigo_matricula)
    else:
        exp = Expediente2.objects.get_or_create(estudiante=estudiante)[0]
        campos_validos = ['certificado_notas', 'constancia_proyeccion', 'constancia_no_adeudo']
        nombres = renombrar_exp2(estudiante.apellidos_nombres)

    if campo not in campos_validos:
        return JsonResponse({'ok': False, 'error': 'Campo inválido'}, status=400)

    # Eliminar archivo anterior si existe
    campo_actual = getattr(exp, campo)
    if campo_actual:
        if os.path.exists(campo_actual.path):
            os.remove(campo_actual.path)

    # Guardar con nombre renombrado
    nombre_final = nombres.get(campo, archivo.name)
    dni = estudiante.dni or 'sin_dni'
    carpeta = f'expedientes/{dni}/exp1/' if expediente_tipo == 'exp1' else f'expedientes/{dni}/exp2/'
    ruta_final = carpeta + nombre_final

    ruta_guardada = default_storage.save(ruta_final, ContentFile(archivo.read()))
    setattr(exp, campo, ruta_guardada)
    exp.save()

    return JsonResponse({
        'ok': True,
        'url': default_storage.url(ruta_guardada),
        'nombre': nombre_final,
        'semaforo': estudiante.estado_semaforo(),
        'archivos': estudiante.archivos_estado(),
    })


# ──────────────────────────────────────────────────────────────
# ELIMINAR ARCHIVO (AJAX)
# ──────────────────────────────────────────────────────────────

@login_required
@require_POST
def eliminar_archivo(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    data = json.loads(request.body)
    expediente_tipo = data.get('expediente')
    campo = data.get('campo')

    if expediente_tipo == 'exp1':
        exp = get_object_or_404(Expediente1, estudiante=estudiante)
    else:
        exp = get_object_or_404(Expediente2, estudiante=estudiante)

    campo_archivo = getattr(exp, campo, None)
    if campo_archivo:
        if os.path.exists(campo_archivo.path):
            os.remove(campo_archivo.path)
        setattr(exp, campo, None)
        exp.save()

    return JsonResponse({
        'ok': True,
        'semaforo': estudiante.estado_semaforo(),
        'archivos': estudiante.archivos_estado(),
    })


# ──────────────────────────────────────────────────────────────
# DESCARGAR EXPEDIENTE 1 (ZIP)
# ──────────────────────────────────────────────────────────────

@login_required
def descargar_expediente1(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    exp1 = get_object_or_404(Expediente1, estudiante=estudiante)

    dni = estudiante.dni or 'SIN_DNI'
    nombre_zip = f'Informe_R06_{dni}_FAC.zip'
    nombres = renombrar_exp1(estudiante.dni, estudiante.apellidos_nombres, estudiante.codigo_matricula)

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        campos = ['dni_pdf', 'solicitud_pdf', 'foto', 'recibo_pdf', 'registro_graduados_pdf']
        for campo in campos:
            archivo = getattr(exp1, campo)
            if archivo:
                try:
                    nombre_final = nombres.get(campo, os.path.basename(archivo.name))
                    with open(archivo.path, 'rb') as f:
                        zf.writestr(nombre_final, f.read())
                except Exception:
                    pass

    buffer.seek(0)
    response = HttpResponse(buffer.read(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{nombre_zip}"'
    return response


# ──────────────────────────────────────────────────────────────
# DESCARGAR EXPEDIENTE 2 (PDF UNIFICADO)
# ──────────────────────────────────────────────────────────────

@login_required
def descargar_expediente2_unido(request, pk):
    """Une los 3 PDFs del expediente 2 en uno solo usando pypdf"""
    estudiante = get_object_or_404(Estudiante, pk=pk)
    exp2 = get_object_or_404(Expediente2, estudiante=estudiante)

    dni = estudiante.dni or 'SIN_DNI'
    nombre_salida = f'R07-FAC-{dni}-2026-PRUEBA.pdf'

    try:
        from pypdf import PdfWriter, PdfReader

        writer = PdfWriter()
        campos = ['certificado_notas', 'constancia_proyeccion', 'constancia_no_adeudo']
        archivos_unidos = 0

        for campo in campos:
            archivo = getattr(exp2, campo)
            if archivo:
                try:
                    reader = PdfReader(archivo.path)
                    for page in reader.pages:
                        writer.add_page(page)
                    archivos_unidos += 1
                except Exception:
                    pass

        if archivos_unidos == 0:
            return JsonResponse({'error': 'No hay archivos PDF para unir'}, status=400)

        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{nombre_salida}"'
        return response

    except ImportError:
        # Fallback: descargar en ZIP si pypdf no está instalado
        buffer = io.BytesIO()
        nombres = renombrar_exp2(estudiante.apellidos_nombres)
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for campo in ['certificado_notas', 'constancia_proyeccion', 'constancia_no_adeudo']:
                archivo = getattr(exp2, campo)
                if archivo:
                    try:
                        nombre_final = nombres.get(campo, os.path.basename(archivo.name))
                        with open(archivo.path, 'rb') as f:
                            zf.writestr(nombre_final, f.read())
                    except Exception:
                        pass
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="R07-FAC-{dni}-2026-PRUEBA.zip"'
        return response


# ──────────────────────────────────────────────────────────────
# GENERAR INFORME WORD
# ──────────────────────────────────────────────────────────────

@login_required
def generar_informe_word(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    fecha_str = request.GET.get('fecha', '')
    numero_informe = request.GET.get('numero_informe', '')

    # Guardar n_informe en BD si se proporcionó
    if numero_informe:
        estudiante.n_informe = numero_informe
        estudiante.save(update_fields=['n_informe'])
    dni = estudiante.dni or 'SIN_DNI'
    n_exp = estudiante.n_expediente or ''
    nombre_archivo = f'INFORME_{n_exp}_{dni}.docx'

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        # Verificar si hay plantilla personalizada
        plantilla_path = settings.WORD_TEMPLATES_DIR / 'plantilla_informe.docx'

        if plantilla_path.exists():
            doc = Document(str(plantilla_path))
            # Reemplazar variables en plantilla
            reemplazos = {
                '{{FECHA}}': fecha_str,
                '{{NOMBRE}}': estudiante.apellidos_nombres or '',
                '{{DNI}}': estudiante.dni or '',
                '{{EXPEDIENTE}}': estudiante.n_expediente or '',
                '{{PROVEIDO}}': estudiante.proveido or '',
                '{{ESCUELA}}': dict(ESCUELAS).get(estudiante.escuela, ''),
                '{{CELULAR}}': estudiante.celular or '',
                '{{CORREO}}': estudiante.correo or '',
                '{{INFORME}}': numero_informe,
                '{{CODIGO_MATRICULA}}': estudiante.codigo_matricula or '',
            }
            for para in doc.paragraphs:
                for var, val in reemplazos.items():
                    if var in para.text:
                        for run in para.runs:
                            run.text = run.text.replace(var, val)
        else:
            # Generar documento desde cero
            doc = Document()

            # Estilo del documento
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

            # Encabezado institucional
            titulo = doc.add_heading('', 0)
            run = titulo.add_run('UNIVERSIDAD NACIONAL')
            run.font.color.rgb = RGBColor(0, 101, 177)
            run.font.size = Pt(14)
            run.bold = True

            subtitulo = doc.add_heading('', 2)
            run2 = subtitulo.add_run('ÁREA DE REGISTROS ACADÉMICOS')
            run2.font.color.rgb = RGBColor(0, 101, 177)
            run2.font.size = Pt(12)

            doc.add_paragraph()

            # Número de informe
            p_informe = doc.add_paragraph()
            p_informe.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_inf = p_informe.add_run(f'INFORME N° {n_exp}')
            run_inf.bold = True
            run_inf.font.size = Pt(13)

            doc.add_paragraph()

            # Fecha
            p_fecha = doc.add_paragraph()
            p_fecha.add_run('Fecha: ').bold = True
            p_fecha.add_run(fecha_str)

            doc.add_paragraph()

            # Datos del expediente en tabla
            tabla = doc.add_table(rows=10, cols=2)
            tabla.style = 'Table Grid'

            datos = [
                ('Apellidos y Nombres:', estudiante.apellidos_nombres or ''),
                ('N° DNI:', estudiante.dni or ''),
                ('N° Expediente:', estudiante.n_expediente or ''),
                ('Proveído:', estudiante.proveido or ''),
                ('Escuela Profesional:', dict(ESCUELAS).get(estudiante.escuela, '')),
                ('Celular:', estudiante.celular or ''),
                ('Correo Electrónico:', estudiante.correo or ''),
                ('Código de Matrícula:', estudiante.codigo_matricula or ''),
                ('N° de Informe:', numero_informe),
                ('Fecha de Registro:', fecha_str),
            ]

            for i, (label, valor) in enumerate(datos):
                row = tabla.rows[i]
                cell_label = row.cells[0]
                cell_valor = row.cells[1]

                cell_label.text = label
                cell_valor.text = valor

                for run in cell_label.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 101, 177)

            doc.add_paragraph()

            # Estado de documentos
            doc.add_heading('Estado de Documentos', level=2)

            archivos_estado = estudiante.archivos_estado()

            doc.add_heading('Expediente 1:', level=3)
            labels_exp1 = {
                'dni_pdf': 'DNI (R05)',
                'solicitud_pdf': 'Solicitud (R04)',
                'foto': 'Fotografía',
                'recibo_pdf': 'Recibo de Pago (R03)',
                'registro_graduados_pdf': 'Registro de Graduados (R01)',
            }
            for campo, label in labels_exp1.items():
                p = doc.add_paragraph(style='List Bullet')
                estado = '✓ ENTREGADO' if archivos_estado['exp1'][campo] else '✗ PENDIENTE'
                p.add_run(f'{label}: ').bold = True
                p.add_run(estado)

            doc.add_heading('Expediente 2:', level=3)
            labels_exp2 = {
                'certificado_notas': 'Certificado de Notas',
                'constancia_proyeccion': 'Constancia de Proyección Social',
                'constancia_no_adeudo': 'Constancia de No Adeudo',
            }
            for campo, label in labels_exp2.items():
                p = doc.add_paragraph(style='List Bullet')
                estado = '✓ ENTREGADO' if archivos_estado['exp2'][campo] else '✗ PENDIENTE'
                p.add_run(f'{label}: ').bold = True
                p.add_run(estado)

            doc.add_paragraph()
            p_firma = doc.add_paragraph()
            p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_firma.add_run('\n\n_______________________________\n').bold = True
            p_firma.add_run('Responsable de Registros Académicos')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
        return response

    except ImportError:
        return JsonResponse({
            'error': 'python-docx no está instalado. Ejecuta: pip install python-docx'
        }, status=500)


# ──────────────────────────────────────────────────────────────
# RESOLUCIONES
# ──────────────────────────────────────────────────────────────

@login_required
@require_POST
def subir_resolucion(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    nombre = request.POST.get('nombre', '').strip()
    archivo = request.FILES.get('archivo')

    if not nombre or not archivo:
        return JsonResponse({'ok': False, 'error': 'Nombre y archivo requeridos'}, status=400)

    resolucion = Resolucion.objects.create(
        estudiante=estudiante,
        nombre=nombre,
        archivo=archivo,
    )

    return JsonResponse({
        'ok': True,
        'id': resolucion.pk,
        'nombre': resolucion.nombre,
        'url': resolucion.archivo.url,
        'fecha': resolucion.fecha_subida.strftime('%d/%m/%Y'),
    })


@login_required
@require_POST
def eliminar_resolucion(request, pk):
    resolucion = get_object_or_404(Resolucion, pk=pk)
    if resolucion.archivo and os.path.exists(resolucion.archivo.path):
        os.remove(resolucion.archivo.path)
    resolucion.delete()
    return JsonResponse({'ok': True})


# ──────────────────────────────────────────────────────────────
# API: ESTADO SEMÁFORO (para actualización en tiempo real)
# ──────────────────────────────────────────────────────────────

@login_required
def api_estado_estudiante(request, pk):
    estudiante = get_object_or_404(Estudiante, pk=pk)
    return JsonResponse({
        'semaforo': estudiante.estado_semaforo(),
        'archivos': estudiante.archivos_estado(),
        'nombre': estudiante.get_nombre_display(),
    })
